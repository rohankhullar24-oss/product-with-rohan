#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Build "The Product Manager Handbook" by merging 93 raw chapter .docx files into
one professionally formatted Word document (and PDF), organized into 5 volumes
with front matter, appendices, glossary, and index.

Usage:
    python3 build_handbook.py

See README.md in this folder for details and regeneration instructions.
"""
import glob
import json
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import content as C

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRATCH = "/tmp/claude-0/-home-user-product-with-rohan/7c38b8a3-ef9f-5ea8-afe0-e0f168c5410f/scratchpad"
SRC_DIR = os.path.join(SCRATCH, "pm-handbook-src")
REPO_ROOT = "/home/user/product-with-rohan"
OUT_DIR = os.path.join(REPO_ROOT, "public", "handbook")
OUT_DOCX = os.path.join(OUT_DIR, "The-Product-Manager-Handbook.docx")
MANIFEST_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "manifest.json")

# ---------------------------------------------------------------------------
# Brand colors / typography (from tailwind.config.ts)
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x0F, 0x17, 0x2A)
SLATE = RGBColor(0x47, 0x55, 0x69)
ACCENT = RGBColor(0x0D, 0x94, 0x88)
BODY_BLACK = RGBColor(0x20, 0x20, 0x20)

HEADING_FONT = "Cambria"
BODY_FONT = "Calibri"
CODE_FONT = "Consolas"

VOLUME_NUM_BY_ROMAN = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5}


# ---------------------------------------------------------------------------
# Chapter discovery / ordering
# ---------------------------------------------------------------------------
def chapter_num(fname):
    m = re.search(r'Chapter_?(\d+)', fname)
    return int(m.group(1))


def volume_of(fname):
    base = os.path.basename(fname)
    if 'Volume2' in base:
        return 2
    if 'Volume3' in base:
        return 3
    if 'Volume4' in base:
        return 4
    if base.startswith('Volume5'):
        return 5
    return 1


def discover_chapters():
    files = sorted(glob.glob(os.path.join(SRC_DIR, "*.docx")))
    files = [f for f in files if 'Master_Manuscript' not in f]
    groups = {1: [], 2: [], 3: [], 4: [], 5: []}
    for f in files:
        groups[volume_of(f)].append(f)
    for v in groups:
        groups[v].sort(key=lambda f: chapter_num(os.path.basename(f)))
    total = sum(len(v) for v in groups.values())
    assert total == 93, f"Expected 93 chapters, found {total}"
    assert [len(groups[v]) for v in (1, 2, 3, 4, 5)] == [16, 29, 16, 16, 16], \
        {v: len(groups[v]) for v in groups}
    return groups


# ---------------------------------------------------------------------------
# Extraction — read each chapter into a neutral block format
# ---------------------------------------------------------------------------
SRC_STYLE_MAP = {
    'Heading 2': 'Heading 1',   # chapter title (handled specially, see below)
    'Heading 3': 'Heading 2',   # subsection
    'Heading 4': 'Heading 3',   # sub-subsection
    'Normal': 'Normal',
    'List Bullet': 'List Bullet',
    'List Bullet 2': 'List Bullet 2',
    'List Bullet 3': 'List Bullet 2',
    'List Number': 'List Number',
    'List Number 2': 'List Number',
    'Intense Quote': 'Intense Quote',
    'Body Text': 'Normal',
}


def extract_chapter(path):
    """Return (chapter_title, [blocks]) using the shared block dict format."""
    doc = Document(path)
    body = doc.element.body
    blocks = []
    skipped_title = False
    skipped_h1 = False
    chapter_title = None

    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            p = Paragraph(child, doc)
            style_name = p.style.name if p.style is not None else 'Normal'
            text = p.text

            if style_name == 'Title' and not skipped_title:
                skipped_title = True
                continue
            if style_name == 'Heading 1' and not skipped_h1:
                skipped_h1 = True
                continue
            if style_name == 'Heading 2' and chapter_title is None:
                chapter_title = text.strip()
                continue

            if not text.strip():
                continue  # drop blank paragraphs; spacing is style-driven

            mapped = SRC_STYLE_MAP.get(style_name, 'Normal')
            runs = [{'text': r.text, 'bold': bool(r.bold)} for r in p.runs if r.text]
            if not runs:
                runs = [{'text': text, 'bold': False}]
            blocks.append({'type': 'p', 'style': mapped, 'runs': runs})

        elif child.tag == qn('w:tbl'):
            t = Table(child, doc)
            rows = [[cell.text for cell in row.cells] for row in t.rows]
            blocks.append({'type': 'tbl', 'rows': rows})

    if chapter_title is None:
        raise ValueError(f"No Heading 2 chapter title found in {path}")
    return chapter_title, blocks


# ---------------------------------------------------------------------------
# New document — style definitions
# ---------------------------------------------------------------------------
def set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def define_styles(doc):
    styles = doc.styles
    existing = [s.name for s in styles]

    def get_or_add(name, style_type=WD_STYLE_TYPE.PARAGRAPH):
        if name in existing:
            return styles[name]
        return styles.add_style(name, style_type)

    # Normal (body) ----------------------------------------------------------
    normal = styles['Normal']
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_BLACK
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)

    # Title / Subtitle ---------------------------------------------------------
    title = styles['Title']
    title.font.name = HEADING_FONT
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = NAVY
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    subtitle = get_or_add('Subtitle')
    subtitle.base_style = normal
    subtitle.font.name = HEADING_FONT
    subtitle.font.size = Pt(16)
    subtitle.font.italic = True
    subtitle.font.color.rgb = SLATE
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)

    # Headings -----------------------------------------------------------------
    h1 = styles['Heading 1']
    h1.font.name = HEADING_FONT
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.keep_with_next = True

    h2 = styles['Heading 2']
    h2.font.name = HEADING_FONT
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    h3 = styles['Heading 3']
    h3.font.name = HEADING_FONT
    h3.font.size = Pt(12.5)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = SLATE
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True

    # Lists ----------------------------------------------------------------
    for lname in ('List Bullet', 'List Number'):
        st = styles[lname]
        st.font.name = BODY_FONT
        st.font.size = Pt(11)
        st.font.color.rgb = BODY_BLACK
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.15

    lb2 = get_or_add('List Bullet 2')
    lb2.base_style = styles['List Bullet']
    lb2.paragraph_format.left_indent = Cm(1.6)

    # Intense Quote -----------------------------------------------------------
    iq = get_or_add('Intense Quote')
    iq.base_style = normal
    iq.font.name = HEADING_FONT
    iq.font.size = Pt(11.5)
    iq.font.bold = True
    iq.font.italic = True
    iq.font.color.rgb = ACCENT
    iq.paragraph_format.space_before = Pt(10)
    iq.paragraph_format.space_after = Pt(6)
    iq.paragraph_format.left_indent = Cm(0.5)

    # Code -------------------------------------------------------------------
    code = get_or_add('Code')
    code.base_style = normal
    code.font.name = CODE_FONT
    code.font.size = Pt(9.5)
    code.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0

    # Caption ------------------------------------------------------------------
    cap = get_or_add('Caption')
    cap.base_style = normal
    cap.font.name = BODY_FONT
    cap.font.size = Pt(9.5)
    cap.font.italic = True
    cap.font.color.rgb = SLATE
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(10)

    # Volume divider ------------------------------------------------------------
    vdiv = get_or_add('VolumeDivider')
    vdiv.base_style = h1
    vdiv.font.size = Pt(26)
    vdiv.font.color.rgb = ACCENT
    vdiv.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vdiv.paragraph_format.space_before = Pt(60)
    vdiv.paragraph_format.space_after = Pt(24)


# ---------------------------------------------------------------------------
# Field helpers (TOC, INDEX, XE, PAGE)
# ---------------------------------------------------------------------------
def add_field(paragraph, instr_text, result_text=None):
    """Insert a Word complex field (begin/instrText/separate/result/end)."""
    run = paragraph.add_run()

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instr_text
    run._r.append(instr)

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)

    result_run = paragraph.add_run(result_text if result_text is not None else '')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    result_run._r.append(fld_end)
    return run


def add_xe_field(paragraph, term):
    """Insert a hidden XE (index-entry) field at the end of a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    safe_term = term.replace('"', "'")
    instr.text = f' XE "{safe_term}" '
    run._r.append(instr)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)


def add_page_field(paragraph):
    add_field(paragraph, ' PAGE ', result_text='1')


def set_update_fields(doc):
    """Force Word/LibreOffice to recompute all fields (TOC, INDEX, PAGE) on open."""
    settings = doc.settings.element
    uf = OxmlElement('w:updateFields')
    uf.set(qn('w:val'), 'true')
    settings.append(uf)


# ---------------------------------------------------------------------------
# Section / header / footer / page-number helpers
# ---------------------------------------------------------------------------
def set_page_number_format(section, fmt='decimal', start=None):
    """Set the page-number format for a section. If start is None, numbering
    continues from the previous section -- this explicitly clears any w:start
    attribute, because python-docx's add_section() clones the previous
    section's sectPr, which would otherwise carry over a stale restart value
    (e.g. every volume silently restarting at page 1 instead of continuing)."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    elif pgNumType.get(qn('w:start')) is not None:
        del pgNumType.attrib[qn('w:start')]


def setup_footer_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p.text = ''
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(p)


def setup_header_text(section, text):
    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        p.text = ''
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = SLATE


def apply_margins(section):
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)


# ---------------------------------------------------------------------------
# Emitting blocks (shared by chapters + appendices + front/back matter)
# ---------------------------------------------------------------------------
# A few index terms are referenced in the source chapters by a shorter form than
# their glossary/display name (e.g. tables say "RICE", the term is "RICE Score").
EXTRA_INDEX_ALTS = {
    'RICE Score': ['RICE'],
    'ICE Score': ['ICE'],
}


def make_index_matchers(terms):
    compiled = []
    for term in terms:
        m = re.match(r'^(.*?)\s*\((.*?)\)\s*$', term)
        alts = [m.group(1).strip(), m.group(2).strip()] if m else [term]
        alts += EXTRA_INDEX_ALTS.get(term, [])
        patterns = []
        for a in alts:
            esc = re.escape(a).replace(r'\ ', r'[\s\-]+')
            patterns.append(re.compile(r'\b' + esc + r'\b', re.IGNORECASE))
        compiled.append({'term': term, 'patterns': patterns, 'done': False})
    return compiled


def emit_blocks(doc, blocks, index_matchers=None, table_seq_ref=None, volume_num=None):
    for block in blocks:
        if block['type'] == 'p':
            p = doc.add_paragraph(style=block['style'])
            full_text = ''.join(r['text'] for r in block['runs'])
            for r in block['runs']:
                run = p.add_run(r['text'])
                if r.get('bold'):
                    run.bold = True
            if index_matchers is not None:
                for entry in index_matchers:
                    if entry['done']:
                        continue
                    if any(pat.search(full_text) for pat in entry['patterns']):
                        add_xe_field(p, entry['term'])
                        entry['done'] = True
        elif block['type'] == 'tbl':
            rows = block['rows']
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                row_text = ' | '.join(row)
                first_cell_para = None
                for ci in range(ncols):
                    text = row[ci] if ci < len(row) else ''
                    cell = table.cell(ri, ci)
                    cell.text = ''
                    para = cell.paragraphs[0]
                    if first_cell_para is None:
                        first_cell_para = para
                    run = para.add_run(text)
                    run.font.name = BODY_FONT
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        set_cell_shading(cell, '0F172A')
                    else:
                        run.font.color.rgb = BODY_BLACK
                        if ri % 2 == 0:
                            set_cell_shading(cell, 'F1F5F9')
                if index_matchers is not None and first_cell_para is not None:
                    for entry in index_matchers:
                        if entry['done']:
                            continue
                        if any(pat.search(row_text) for pat in entry['patterns']):
                            add_xe_field(first_cell_para, entry['term'])
                            entry['done'] = True
            if table_seq_ref is not None and volume_num is not None:
                table_seq_ref[volume_num] = table_seq_ref.get(volume_num, 0) + 1
                cap = doc.add_paragraph(style='Caption')
                cap.add_run(f"Table {volume_num}.{table_seq_ref[volume_num]}")
            spacer = doc.add_paragraph(style='Normal')
            spacer.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Front matter
# ---------------------------------------------------------------------------
def build_front_matter(doc):
    p = doc.add_paragraph(style='Title')
    p.add_run(C.BOOK_TITLE)
    sub = doc.add_paragraph(style='Subtitle')
    sub.add_run(C.BOOK_SUBTITLE)
    for _ in range(6):
        doc.add_paragraph(style='Normal')
    byline = doc.add_paragraph(style='Normal')
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = byline.add_run(C.BOOK_BYLINE)
    r.bold = True
    r.font.size = Pt(15)
    r.font.name = HEADING_FONT
    r.font.color.rgb = NAVY
    ed = doc.add_paragraph(style='Normal')
    ed.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = ed.add_run(C.BOOK_EDITION)
    er.font.size = Pt(12)
    er.font.color.rgb = SLATE
    er.italic = True

    # Copyright page
    first = True
    for block in C.COPYRIGHT_BLOCKS:
        p2 = doc.add_paragraph(style=block['style'])
        if first:
            p2.paragraph_format.page_break_before = True
            first = False
        for r in block['runs']:
            run = p2.add_run(r['text'])
            if r.get('bold'):
                run.bold = True

    # Dedication
    ded_first = True
    for block in C.DEDICATION:
        p3 = doc.add_paragraph(style='Normal')
        if ded_first:
            p3.paragraph_format.page_break_before = True
            ded_first = False
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(4)
        for r in block['runs']:
            run = p3.add_run(r['text'])
            run.italic = True
            run.font.size = Pt(13)

    # Foreword
    fw_title = doc.add_paragraph(style='Heading 1')
    fw_title.paragraph_format.page_break_before = True
    fw_title.add_run('Foreword')
    emit_blocks(doc, C.FOREWORD_BLOCKS)

    # Preface
    pre_title = doc.add_paragraph(style='Heading 1')
    pre_title.paragraph_format.page_break_before = True
    pre_title.add_run('Preface')
    emit_blocks(doc, C.PREFACE_BLOCKS)

    # How to Use This Book
    htu_title = doc.add_paragraph(style='Heading 1')
    htu_title.paragraph_format.page_break_before = True
    htu_title.add_run('How to Use This Book')
    emit_blocks(doc, C.HOW_TO_USE_BLOCKS)

    # About the Author
    aa_title = doc.add_paragraph(style='Heading 1')
    aa_title.paragraph_format.page_break_before = True
    aa_title.add_run('About the Author')
    emit_blocks(doc, C.ABOUT_AUTHOR_BLOCKS)

    # Table of Contents
    toc_title = doc.add_paragraph(style='Heading 1')
    toc_title.paragraph_format.page_break_before = True
    toc_title.add_run('Table of Contents')
    toc_para = doc.add_paragraph(style='Normal')
    add_field(toc_para, r'TOC \o "1-3" \h \z \u',
              result_text='(Table of Contents will be generated here when the field updates.)')


# ---------------------------------------------------------------------------
# Assemble the whole document
# ---------------------------------------------------------------------------
def build_document():
    groups = discover_chapters()
    doc = Document()
    define_styles(doc)

    front_section = doc.sections[0]
    apply_margins(front_section)
    setup_header_text(front_section, C.BOOK_TITLE)
    setup_footer_page_number(front_section)
    set_page_number_format(front_section, fmt='lowerRoman', start=1)

    build_front_matter(doc)

    index_matchers = make_index_matchers(C.INDEX_TERMS)
    table_seq = {}
    manifest_volumes = []

    for i, vol_meta in enumerate(C.VOLUME_META):
        vol_num = VOLUME_NUM_BY_ROMAN[vol_meta['roman']]
        header_text = f"Volume {vol_meta['roman']} — {vol_meta['title']}"

        sect = doc.add_section(WD_SECTION.NEW_PAGE)
        apply_margins(sect)
        setup_header_text(sect, header_text)
        setup_footer_page_number(sect)
        if i == 0:
            set_page_number_format(sect, fmt='decimal', start=1)
        else:
            set_page_number_format(sect, fmt='decimal')  # continue numbering

        # Volume divider (new section already forced a page break; no extra one needed)
        divider = doc.add_paragraph(style='VolumeDivider')
        divider.add_run(f"VOLUME {vol_meta['roman']} — {vol_meta['title'].upper()}")
        emit_blocks(doc, vol_meta['intro'])

        chapter_titles = []
        for path in groups[vol_num]:
            chapter_title, blocks = extract_chapter(path)
            chapter_titles.append(chapter_title)

            h = doc.add_paragraph(style='Heading 1')
            h.paragraph_format.page_break_before = True
            h.add_run(chapter_title)
            for entry in index_matchers:
                if entry['done']:
                    continue
                if any(pat.search(chapter_title) for pat in entry['patterns']):
                    add_xe_field(h, entry['term'])
                    entry['done'] = True

            emit_blocks(doc, blocks, index_matchers=index_matchers,
                        table_seq_ref=table_seq, volume_num=vol_num)

        manifest_volumes.append({
            'volume': f"Volume {vol_meta['roman']}",
            'title': vol_meta['title'],
            'description': vol_meta['short_desc'],
            'chapters': chapter_titles,
        })

    # Back matter section
    back_sect = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_margins(back_sect)
    setup_header_text(back_sect, "Appendices, Glossary & Index")
    setup_footer_page_number(back_sect)
    set_page_number_format(back_sect, fmt='decimal')  # continue numbering

    appendix_titles = build_back_matter(doc, index_matchers)

    set_update_fields(doc)

    return doc, manifest_volumes, appendix_titles, index_matchers


def build_back_matter(doc, index_matchers):
    appendix_titles = []

    appendices_intro = doc.add_paragraph(style='Heading 1')
    appendices_intro.add_run('Appendices')
    intro_p = doc.add_paragraph(style='Normal')
    intro_p.add_run("A compact set of reference materials to use alongside the five volumes "
                     "— frameworks, templates, a KPI library, SQL patterns, an AI prompt "
                     "library, and interview and resume guidance.")

    for builder in (C.build_appendix_a, C.build_appendix_b, C.build_appendix_c,
                    C.build_appendix_d, C.build_appendix_e, C.build_appendix_f,
                    C.build_appendix_g, C.build_appendix_h):
        blocks = builder()
        title_block = blocks[0]
        title_text = ''.join(r['text'] for r in title_block['runs'])
        appendix_titles.append(title_text)
        h = doc.add_paragraph(style='Heading 1')
        h.paragraph_format.page_break_before = True
        h.add_run(title_text)
        emit_blocks(doc, blocks[1:])

    # Glossary
    gloss_title = doc.add_paragraph(style='Heading 1')
    gloss_title.paragraph_format.page_break_before = True
    gloss_title.add_run('Glossary')
    doc.add_paragraph(style='Normal').add_run(
        "Key terms and frameworks used throughout this handbook, alphabetized for quick lookup."
    )
    for term, definition in C.GLOSSARY_TERMS:
        p = doc.add_paragraph(style='Normal')
        r1 = p.add_run(term + '. ')
        r1.bold = True
        p.add_run(definition)

    # Index
    idx_title = doc.add_paragraph(style='Heading 1')
    idx_title.paragraph_format.page_break_before = True
    idx_title.add_run('Index')
    note = doc.add_paragraph(style='Normal')
    note.add_run("This index is generated from index-entry fields placed at the first "
                 "occurrence of each term in the body of the book. If page numbers below "
                 "look out of date, select this document (Ctrl+A) and press F9 to refresh.")
    idx_para = doc.add_paragraph(style='Normal')
    add_field(idx_para, r'INDEX \e "  " \h "A" \z "1033"',
              result_text='(Index will be generated here when the field updates.)')

    return appendix_titles


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc, manifest_volumes, appendix_titles, index_matchers = build_document()
    doc.save(OUT_DOCX)

    unmatched = [e['term'] for e in index_matchers if not e['done']]
    print(f"Saved {OUT_DOCX}")
    print(f"Index terms placed: {len(index_matchers) - len(unmatched)}/{len(index_matchers)}")
    if unmatched:
        print("WARNING: index terms never matched in body text:", unmatched)

    manifest = {
        'title': C.BOOK_TITLE,
        'subtitle': C.BOOK_SUBTITLE,
        'byline': C.BOOK_BYLINE,
        'edition': C.BOOK_EDITION,
        'volumes': manifest_volumes,
        'appendices': appendix_titles,
        'docx_path': OUT_DOCX,
    }
    with open(MANIFEST_PATH, 'w') as f:
        json.dump(manifest, f, indent=2, ensure_ascii=False)
    print(f"Wrote manifest skeleton to {MANIFEST_PATH} (page counts/file sizes added after PDF conversion)")


if __name__ == '__main__':
    main()
